"""
Streamlit GUI for Travel Planner
"""

import streamlit as st
from datetime import datetime, timedelta
import json
import os
import re
import time
import uuid
from dotenv import load_dotenv
from openai import OpenAI
import pandas as pd
from io import BytesIO

# Load environment variables
load_dotenv()

# Import orchestrator
from overarching_bot import plan_trip, Strategy, parse_user_input
from airline_codes import get_airline_with_code, resolve_airline_code
from weather import get_weather_forecast
from events_bot import search_events

# Initialize OpenAI client
client = OpenAI(timeout=30.0)

# Constants
MAX_CHAT_HISTORY = 50
API_TIMEOUT = 30
RATE_LIMIT_SECONDS = 2
MAX_TRIP_HISTORY = 20

# Page config
st.set_page_config(
    page_title="AI Travel Planner",
    page_icon="✈️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ==========================================================
# CUSTOM CSS
# ==========================================================
st.markdown("""
<style>
    .main-header {
        background: linear-gradient(135deg, #FF6B6B, #4ECDC4);
        padding: 2rem;
        border-radius: 15px;
        margin-bottom: 2rem;
        text-align: center;
    }
    .main-header h1 {
        color: #FFFFFF !important;
    }
    .chat-message-user {
        background: linear-gradient(135deg, #FF6B6B20, #4ECDC420);
        padding: 1rem;
        border-radius: 20px 20px 5px 20px;
        border: 2px solid #FF6B6B;
        margin: 0.5rem 0;
        text-align: right;
    }
    .chat-message-bot {
        background: #F8F9FA;
        padding: 1rem;
        border-radius: 20px 20px 20px 5px;
        border: 2px solid #4ECDC4;
        margin: 0.5rem 0;
    }
    .weather-card {
        background: linear-gradient(135deg, #4ECDC4, #45B7AA);
        padding: 1rem;
        border-radius: 10px;
        margin: 1rem 0;
        color: white;
    }
    .red-eye-badge {
        background-color: #FF6B6B;
        color: white;
        padding: 0.2rem 0.5rem;
        border-radius: 20px;
        font-size: 0.8rem;
        display: inline-block;
    }
    .events-container {
        max-height: 400px;
        overflow-y: auto;
        padding: 10px;
        background: #f8f9fa;
        border-radius: 10px;
        margin: 10px 0;
        border: 1px solid #4ECDC4;
    }
    .event-day {
        margin-bottom: 15px;
        padding: 10px;
        background: white;
        border-radius: 8px;
        border-left: 3px solid #4ECDC4;
    }
    .event-item {
        padding: 5px 0;
        border-bottom: 1px solid #eee;
        font-size: 0.9rem;
    }
    .event-item:last-child {
        border-bottom: none;
    }
    
    /* Trip History Styles */
    .history-item {
        background: white;
        padding: 12px;
        border-radius: 8px;
        margin: 8px 0;
        border-left: 4px solid #4ECDC4;
        cursor: pointer;
        transition: all 0.2s;
    }
    .history-item:hover {
        transform: translateX(5px);
        border-left: 4px solid #FF6B6B;
        background: #f0f0f0;
    }
    .history-item.current {
        border-left: 4px solid #FF6B6B;
        background: #fff0f0;
    }
    
    /* Trip card button styling */
    .stButton > button[key*="trip_"] {
        white-space: normal !important;
        word-wrap: break-word !important;
        text-align: left !important;
        height: auto !important;
        min-height: 100px !important;
        padding: 15px !important;
        background: white !important;
        border: 2px solid #4ECDC4 !important;
        border-radius: 10px !important;
        color: #2D3436 !important;
        font-weight: normal !important;
        line-height: 1.5 !important;
        box-shadow: 0 2px 5px rgba(0,0,0,0.1) !important;
        transition: all 0.2s !important;
    }
    
    .stButton > button[key*="trip_"]:hover {
        border-color: #FF6B6B !important;
        transform: translateY(-2px) !important;
        box-shadow: 0 4px 10px rgba(0,0,0,0.15) !important;
    }
    
    .stButton > button[key*="trip_"] strong {
        font-size: 1.1rem !important;
        color: #FF6B6B !important;
        display: block !important;
        margin-bottom: 5px !important;
    }
    
    .stButton > button[key*="trip_"] small {
        font-size: 0.9rem !important;
        color: #666 !important;
        display: block !important;
    }
    
    .stButton > button[key*="trip_"] small:last-child {
        color: #4ECDC4 !important;
        font-weight: bold !important;
        margin-top: 5px !important;
    }
    
    /* Delete button styling - PERFECT CENTERING */
    .stButton > button[key*="del_"] {
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
        height: 100px !important;
        width: 100% !important;
        padding: 0 !important;
        margin: 0 !important;
        background: white !important;
        border: 2px solid #FF6B6B !important;
        border-radius: 10px !important;
        font-size: 2rem !important;
        line-height: 1 !important;
        color: #FF6B6B !important;
        transition: all 0.2s !important;
        box-shadow: 0 2px 5px rgba(0,0,0,0.1) !important;
    }
    
    .stButton > button[key*="del_"]:hover {
        background: #FF6B6B !important;
        color: white !important;
        transform: translateY(-2px) !important;
        box-shadow: 0 4px 10px rgba(0,0,0,0.15) !important;
        border-color: #FF6B6B !important;
    }
    
    /* Save button styling */
    div.stDownloadButton > button {
        background: linear-gradient(135deg, #FF6B6B, #4ECDC4);
        color: white;
        border: none;
        border-radius: 25px;
        padding: 0.5rem 2rem;
        font-weight: 600;
        width: 100%;
    }
    div.stDownloadButton > button:hover {
        opacity: 0.9;
    }
    
    /* General button styling for other buttons */
    .stButton > button:not([key*="trip_"]):not([key*="del_"]) {
        background: linear-gradient(135deg, #FF6B6B, #4ECDC4);
        color: white;
        border: none;
        border-radius: 25px;
        padding: 0.5rem 1rem;
        font-weight: 600;
        transition: opacity 0.2s;
    }
    .stButton > button:not([key*="trip_"]):not([key*="del_"]):hover {
        opacity: 0.9;
    }
</style>
""", unsafe_allow_html=True)


# ==========================================================
# TRIP HISTORY FUNCTIONS
# ==========================================================

def save_trip_to_history():
    """Save current trip to history - auto-save on any change"""
    if not st.session_state.trip_result:
        return
    
    # Generate unique ID if new trip
    if st.session_state.current_trip_id == -1:
        st.session_state.current_trip_id = str(uuid.uuid4())[:8]
    
    # Find if this trip already exists
    existing_index = None
    for i, trip in enumerate(st.session_state.trip_history):
        if trip.get("id") == st.session_state.current_trip_id:
            existing_index = i
            break
    
    # Create trip record with complete state
    trip_record = {
        "id": st.session_state.current_trip_id,
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "origin": st.session_state.current_params['origin'],
        "destination": st.session_state.current_params['destination'],
        "dates": f"{st.session_state.current_params['departure_date']} to {st.session_state.current_params['return_date']}",
        "total_cost": st.session_state.trip_result.get('data', {}).get('total_cost', 0),
        "result": st.session_state.trip_result,
        "params": st.session_state.current_params.copy(),
        "chat_history": st.session_state.chat_history.copy(),
        "events": st.session_state.current_events,
        "events_data": st.session_state.current_events_data,
        "excluded_ids": st.session_state.excluded_event_ids.copy(),
        "optimization_iterations": st.session_state.optimization_iterations.copy(),
        "weather_cache": st.session_state.weather_cache.copy(),
    }
    
    # Update or insert
    if existing_index is not None:
        st.session_state.trip_history[existing_index] = trip_record
    else:
        st.session_state.trip_history.insert(0, trip_record)
    
    # Keep only last MAX_TRIP_HISTORY
    if len(st.session_state.trip_history) > MAX_TRIP_HISTORY:
        st.session_state.trip_history = st.session_state.trip_history[:MAX_TRIP_HISTORY]


def load_trip_from_history(trip_id):
    """Load a complete trip from history - restores all state"""
    for trip in st.session_state.trip_history:
        if trip["id"] == trip_id:
            # Restore all state
            st.session_state.trip_result = trip["result"]
            st.session_state.current_params = trip["params"].copy()
            st.session_state.chat_history = trip["chat_history"].copy()
            st.session_state.current_events = trip.get("events")
            st.session_state.current_events_data = trip.get("events_data")
            st.session_state.excluded_event_ids = trip.get("excluded_ids", []).copy()
            st.session_state.optimization_iterations = trip.get("optimization_iterations", []).copy()
            st.session_state.weather_cache = trip.get("weather_cache", {}).copy()
            st.session_state.current_trip_id = trip["id"]
            sync_sidebar_from_current_params()
            st.session_state.trip_version += 1
            st.rerun()
            return
    
    st.error("Trip not found in history")


def delete_trip_from_history(trip_id):
    """Delete a trip from history"""
    st.session_state.trip_history = [t for t in st.session_state.trip_history if t["id"] != trip_id]
    
    # If current trip was deleted, clear it
    if st.session_state.current_trip_id == trip_id:
        st.session_state.trip_result = None
        st.session_state.chat_history = []
        st.session_state.current_events = None
        st.session_state.current_events_data = None
        st.session_state.excluded_event_ids = []
        st.session_state.optimization_iterations = []
        st.session_state.weather_cache = {}
        st.session_state.current_trip_id = -1
        sync_sidebar_from_current_params()
    
    st.rerun()


def sync_sidebar_from_current_params():
    """Keep sidebar widget state aligned with current_params before widgets render."""
    st.session_state.sidebar_origin = st.session_state.current_params['origin']
    st.session_state.sidebar_destination = st.session_state.current_params['destination']
    st.session_state.sidebar_departure = datetime.strptime(
        st.session_state.current_params['departure_date'], "%Y-%m-%d"
    ).date()
    st.session_state.sidebar_return_date = datetime.strptime(
        st.session_state.current_params['return_date'], "%Y-%m-%d"
    ).date()
    st.session_state.sidebar_total_budget = float(st.session_state.current_params['total_budget'])
    st.session_state.sidebar_strategy = st.session_state.current_params['strategy']
    st.session_state.sidebar_prefer_red_eyes = st.session_state.current_params['prefer_red_eyes']
    st.session_state.sidebar_max_iterations = st.session_state.current_params['max_iterations']


def start_new_trip():
    """Start a fresh new trip"""
    st.session_state.trip_result = None
    st.session_state.chat_history = []
    st.session_state.current_events = None
    st.session_state.current_events_data = None
    st.session_state.excluded_event_ids = []
    st.session_state.optimization_iterations = []
    st.session_state.weather_cache = {}
    st.session_state.current_trip_id = -1
    st.session_state.processing_message = False
    sync_sidebar_from_current_params()
    st.rerun()


# ==========================================================
# CHAT FUNCTION SCHEMAS
# ==========================================================

CHAT_FUNCTIONS = [
    {
        "type": "function",
        "function": {
            "name": "search_flights",
            "description": "Search for new flight options",
            "parameters": {
                "type": "object",
                "properties": {
                    "origin": {"type": "string", "description": "Origin city or airport code"},
                    "destination": {"type": "string", "description": "Destination city or airport code"},
                    "departure_date": {"type": "string", "description": "Departure date (YYYY-MM-DD)"},
                    "return_date": {"type": "string", "description": "Return date (YYYY-MM-DD)"}
                },
                "required": ["origin", "destination", "departure_date", "return_date"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "search_hotels",
            "description": "Search for new hotel options",
            "parameters": {
                "type": "object",
                "properties": {
                    "destination": {"type": "string", "description": "Destination city or area"},
                    "check_in": {"type": "string", "description": "Check-in date (YYYY-MM-DD)"},
                    "check_out": {"type": "string", "description": "Check-out date (YYYY-MM-DD)"}
                },
                "required": ["destination", "check_in", "check_out"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "adjust_budget",
            "description": "Change the trip budget",
            "parameters": {
                "type": "object",
                "properties": {
                    "new_budget": {"type": "number", "description": "New budget amount in USD"}
                },
                "required": ["new_budget"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "change_strategy",
            "description": "Change the budget allocation strategy",
            "parameters": {
                "type": "object",
                "properties": {
                    "new_strategy": {
                        "type": "string",
                        "enum": ["cheapest_overall", "splurge_flight", "splurge_hotel"],
                        "description": "New strategy to use"
                    }
                },
                "required": ["new_strategy"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "adjust_iterations",
            "description": "Change the number of optimization iterations",
            "parameters": {
                "type": "object",
                "properties": {
                    "new_iterations": {
                        "type": "integer",
                        "minimum": 1,
                        "maximum": 20,
                        "description": "Number of optimization attempts"
                    }
                },
                "required": ["new_iterations"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "update_dates",
            "description": "Change travel dates",
            "parameters": {
                "type": "object",
                "properties": {
                    "departure_date": {"type": "string", "description": "New departure date (YYYY-MM-DD)"},
                    "return_date": {"type": "string", "description": "New return date (YYYY-MM-DD)"}
                },
                "required": ["departure_date", "return_date"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "update_preferences",
            "description": "Update travel preferences like red-eye flights",
            "parameters": {
                "type": "object",
                "properties": {
                    "prefer_red_eyes": {
                        "type": "boolean",
                        "description": "Whether to prefer red-eye flights (overnight flights between 9PM-5AM)"
                    }
                },
                "required": ["prefer_red_eyes"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "find_better_hotel",
            "description": "Find a cheaper, more expensive, or different hotel than the current one. Re-search alternatives based on price, not star rating. Use this when the user asks for a cheaper hotel, a more expensive hotel, another hotel, a different hotel, or a better-value hotel.",
            "parameters": {
                "type": "object",
                "properties": {
                    "preference": {
                        "type": "string",
                        "enum": ["cheaper", "more_expensive", "different", "upgrade", "any"],
                        "description": "What kind of hotel to look for: cheaper, more_expensive, different, upgrade, or any. upgrade/any are treated as broad price-based re-searches, not higher-star searches."
                    }
                },
                "required": ["preference"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "search_events",
            "description": "Search for local events, concerts, festivals, sports, food events, or things to do during your trip",
            "parameters": {
                "type": "object",
                "properties": {
                    "preferences": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "What kind of events? Examples: ['music'], ['sports', 'food'], ['family', 'outdoor'], ['concerts', 'festivals']",
                    }
                },
                "required": [],
            },
        }
    }  
]

BETTER_HOTEL_KEYWORDS = [
    "find a better hotel",
    "better hotel",
    "upgrade hotel",
    "different hotel",
    "switch hotel",
    "another hotel",
    "nicer hotel",
    "more expensive hotel",
    "find me a more expensive hotel",
    "find a more expensive hotel",
    "higher price hotel",
    "more expensive option",
    "cheaper hotel",
    "find me a cheaper hotel",
    "find a cheaper hotel",
    "cheaper option",
    "hotel under budget",
]

SEARCH_HOTEL_KEYWORDS = [
    "search hotels",
    "show hotels",
    "more hotels",
    "other hotels",
    "hotel options",
    "find hotels",
]

SEARCH_FLIGHT_KEYWORDS = [
    "search flights",
    "show flights",
    "more flights",
    "other flights",
    "flight options",
    "find flights",
]

def _message_has_keyword(user_message: str, keywords: list[str]) -> bool:
    msg = (user_message or "").lower()
    return any(k in msg for k in keywords)

def should_find_better_hotel(user_message: str) -> bool:
    return _message_has_keyword(user_message, BETTER_HOTEL_KEYWORDS)

def should_search_hotels(user_message: str) -> bool:
    return _message_has_keyword(user_message, SEARCH_HOTEL_KEYWORDS)

def should_search_flights(user_message: str) -> bool:
    return _message_has_keyword(user_message, SEARCH_FLIGHT_KEYWORDS)

# ==========================================================
# EXPORT FUNCTIONS
# ==========================================================

def generate_trip_summary_text(trip_result, params, events_data=None, weather_info=None):
    """Generate plain text summary of trip including events"""
    data = trip_result.get("data", {})
    flights = data.get("flights", [])
    hotel = data.get("hotel", {})
    metadata = trip_result.get("metadata", {})
    
    lines = []
    lines.append("=" * 60)
    lines.append("AI TRAVEL PLANNER - TRIP SUMMARY")
    lines.append("=" * 60)
    lines.append("")
    lines.append(f"Trip: {params['origin']} → {params['destination']}")
    lines.append(f"Dates: {params['departure_date']} to {params['return_date']}")
    lines.append(f"Budget: ${params['total_budget']:.2f}")
    lines.append(f"Strategy: {params['strategy']}")
    lines.append("")
    lines.append("-" * 40)
    lines.append("COST SUMMARY")
    lines.append("-" * 40)
    lines.append(f"Total Cost: ${data.get('total_cost', 0):.2f}")
    lines.append(f"Remaining Budget: ${data.get('remaining_budget', 0):.2f}")
    if data.get('flight_ratio'):
        lines.append(f"Allocation: {data['flight_ratio']*100:.0f}% flights / {data['hotel_ratio']*100:.0f}% hotels")
    lines.append("")
    
    lines.append("-" * 40)
    lines.append("FLIGHTS")
    lines.append("-" * 40)
    if flights:
        for i, flight in enumerate(flights, 1):
            lines.append(f"Flight {i}: {flight.get('airline', 'Unknown')} {flight.get('flight_number', '')}")
            lines.append(f"  From: {flight.get('home_airport')} → {flight.get('destination')}")
            lines.append(f"  Depart: {flight.get('departure_date', 'N/A')}")
            lines.append(f"  Arrive: {flight.get('arrival_date', 'N/A')}")
            lines.append(f"  Duration: {flight.get('duration', 'N/A')}")
            lines.append(f"  Cost: ${flight.get('cost', 0):.2f}")
            lines.append("")
    else:
        lines.append("No flights booked")
        lines.append("")
    
    lines.append("-" * 40)
    lines.append("HOTEL")
    lines.append("-" * 40)
    if hotel:
        lines.append(f"Name: {hotel.get('name', 'Unknown')}")
        if hotel.get('rating'):
            lines.append(f"Rating: {hotel.get('rating')}⭐")
        lines.append(f"Cost: ${hotel.get('total', 0):.2f}")
        lines.append(f"Check-in: {params['departure_date']}")
        lines.append(f"Check-out: {params['return_date']}")
    else:
        lines.append("No hotel booked")
    lines.append("")
    
    # EVENTS SECTION
    if events_data and events_data.get("events"):
        lines.append("-" * 40)
        lines.append("LOCAL EVENTS")
        lines.append("-" * 40)
        
        # Group events by date
        events_by_date = {}
        for event in events_data.get("events", []):
            date_key = event.get("start_date", "unknown")
            if date_key not in events_by_date:
                events_by_date[date_key] = []
            events_by_date[date_key].append(event)
        
        # Sort dates
        sorted_dates = sorted([d for d in events_by_date.keys() if d != "unknown"])
        
        for date_key in sorted_dates:
            try:
                dt = datetime.fromisoformat(date_key)
                formatted_date = dt.strftime("%A, %B %d")
            except:
                formatted_date = date_key
            
            lines.append(f"\n{formatted_date}:")
            for event in events_by_date[date_key]:
                title = event.get("title", "Unknown Event")
                time_str = event.get("start_time", "")
                venue = event.get("venue", "")
                
                event_line = f"  • {title}"
                if time_str and time_str != "All day":
                    event_line += f" at {time_str}"
                if venue:
                    event_line += f" @ {venue}"
                lines.append(event_line)
        
        lines.append("")
    
    # WEATHER SECTION
    if weather_info:
        lines.append("-" * 40)
        lines.append("WEATHER FORECAST")
        lines.append("-" * 40)
        lines.append(f"Temperature: {weather_info.get('temperature', 'N/A')}°F")
        lines.append(f"Conditions: {weather_info.get('description', 'N/A')}")
        if weather_info.get('humidity'):
            lines.append(f"Humidity: {weather_info['humidity']}%")
        if weather_info.get('wind'):
            lines.append(f"Wind: {weather_info['wind']} mph")
        lines.append("")
    
    lines.append("=" * 60)
    lines.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    lines.append("=" * 60)
    
    return "\n".join(lines)


def generate_excel_export(trip_result, params, events_data=None, weather_info=None):
    """Generate Excel file with multiple sheets including events and weather"""
    data = trip_result.get("data", {})
    flights = data.get("flights", [])
    hotel = data.get("hotel", {})
    metadata = trip_result.get("metadata", {})
    
    # Create Excel file in memory
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        # Summary sheet
        summary_data = {
            "Item": ["Origin", "Destination", "Departure", "Return", "Budget", "Strategy", 
                    "Total Cost", "Remaining Budget"],
            "Value": [
                params['origin'],
                params['destination'],
                params['departure_date'],
                params['return_date'],
                f"${params['total_budget']:.2f}",
                params['strategy'],
                f"${data.get('total_cost', 0):.2f}",
                f"${data.get('remaining_budget', 0):.2f}"
            ]
        }
        pd.DataFrame(summary_data).to_excel(writer, sheet_name="Summary", index=False)
        
        # Flights sheet
        if flights:
            flights_data = []
            for i, flight in enumerate(flights, 1):
                flights_data.append({
                    "Flight #": i,
                    "Airline": flight.get('airline', 'Unknown'),
                    "Flight Number": flight.get('flight_number', ''),
                    "From": flight.get('home_airport'),
                    "To": flight.get('destination'),
                    "Departure": flight.get('departure_date'),
                    "Arrival": flight.get('arrival_date'),
                    "Duration": flight.get('duration'),
                    "Cost": flight.get('cost', 0)
                })
            pd.DataFrame(flights_data).to_excel(writer, sheet_name="Flights", index=False)
        else:
            pd.DataFrame(["No flights"]).to_excel(writer, sheet_name="Flights", index=False)
        
        # Hotel sheet
        if hotel:
            hotel_data = {
                "Name": [hotel.get('name', 'Unknown')],
                "Rating": [hotel.get('rating', 'N/A')],
                "Cost": [hotel.get('total', 0)],
                "Check-in": [params['departure_date']],
                "Check-out": [params['return_date']]
            }
            pd.DataFrame(hotel_data).to_excel(writer, sheet_name="Hotel", index=False)
        else:
            pd.DataFrame(["No hotel"]).to_excel(writer, sheet_name="Hotel", index=False)
        
        # EVENTS SHEET
        if events_data and events_data.get("events"):
            events_list = []
            for event in events_data.get("events", []):
                events_list.append({
                    "Date": event.get("start_date", ""),
                    "Time": event.get("start_time", ""),
                    "Event": event.get("title", ""),
                    "Category": event.get("category_name", event.get("category", "")),
                    "Venue": event.get("venue", ""),
                    "Description": event.get("description", "")[:100] if event.get("description") else "",
                })
            pd.DataFrame(events_list).to_excel(writer, sheet_name="Events", index=False)
        
        # Weather sheet
        if weather_info:
            weather_data = {
                "Metric": ["Temperature", "Conditions", "Humidity", "Wind"],
                "Value": [
                    f"{weather_info.get('temperature', 'N/A')}°F",
                    weather_info.get('description', 'N/A'),
                    f"{weather_info.get('humidity', 'N/A')}%",
                    f"{weather_info.get('wind', 'N/A')} mph"
                ]
            }
            pd.DataFrame(weather_data).to_excel(writer, sheet_name="Weather", index=False)
    
    output.seek(0)
    return output


def generate_csv_export(trip_result, params, events_data=None, weather_info=None):
    """Generate CSV export"""
    data = trip_result.get("data", {})
    flights = data.get("flights", [])
    hotel = data.get("hotel", {})
    
    lines = ["Type,Detail,Value"]
    lines.append(f"Trip,Origin,{params['origin']}")
    lines.append(f"Trip,Destination,{params['destination']}")
    lines.append(f"Trip,Dates,{params['departure_date']} to {params['return_date']}")
    lines.append(f"Budget,Total,${params['total_budget']:.2f}")
    lines.append(f"Cost,Total,${data.get('total_cost', 0):.2f}")
    
    for i, flight in enumerate(flights):
        lines.append(f"Flight {i+1},Airline,{flight.get('airline')}")
        lines.append(f"Flight {i+1},Number,{flight.get('flight_number')}")
        lines.append(f"Flight {i+1},Cost,${flight.get('cost', 0):.2f}")
    
    if hotel:
        lines.append(f"Hotel,Name,{hotel.get('name')}")
        lines.append(f"Hotel,Cost,${hotel.get('total', 0):.2f}")
    
    # Add events
    if events_data and events_data.get("events"):
        lines.append("")
        lines.append("EVENTS")
        for event in events_data.get("events", []):
            lines.append(f"Event,{event.get('start_date')},{event.get('title')}")
    
    # Add weather
    if weather_info:
        lines.append("")
        lines.append("WEATHER")
        lines.append(f"Weather,Temperature,{weather_info.get('temperature', 'N/A')}°F")
        lines.append(f"Weather,Conditions,{weather_info.get('description', 'N/A')}")
    
    return "\n".join(lines)

def generate_word_export(trip_result, params, events_data=None, weather_info=None):
    """Generate Word document export"""
    try:
        from docx import Document
        from docx.shared import Inches
        doc = Document()
        
        data = trip_result.get("data", {})
        flights = data.get("flights", [])
        hotel = data.get("hotel", {})
        metadata = trip_result.get("metadata", {})
        
        # Title
        doc.add_heading(f'Trip Plan: {params["origin"]} → {params["destination"]}', 0)
        
        # Trip Details
        doc.add_heading('Trip Details', level=1)
        doc.add_paragraph(f'Dates: {params["departure_date"]} to {params["return_date"]}')
        doc.add_paragraph(f'Budget: ${params["total_budget"]:.2f}')
        doc.add_paragraph(f'Strategy: {params["strategy"]}')
        
        # Cost Summary
        doc.add_heading('Cost Summary', level=1)
        doc.add_paragraph(f'Total Cost: ${data.get("total_cost", 0):.2f}')
        doc.add_paragraph(f'Remaining Budget: ${data.get("remaining_budget", 0):.2f}')
        
        # Flights
        doc.add_heading('Flights', level=1)
        if flights:
            for i, flight in enumerate(flights, 1):
                doc.add_heading(f'Flight {i}', level=2)
                doc.add_paragraph(f'Airline: {flight.get("airline", "Unknown")} {flight.get("flight_number", "")}')
                doc.add_paragraph(f'From: {flight.get("home_airport")} → {flight.get("destination")}')
                doc.add_paragraph(f'Depart: {flight.get("departure_date", "N/A")}')
                doc.add_paragraph(f'Arrive: {flight.get("arrival_date", "N/A")}')
                doc.add_paragraph(f'Duration: {flight.get("duration", "N/A")}')
                doc.add_paragraph(f'Cost: ${flight.get("cost", 0):.2f}')
        else:
            doc.add_paragraph('No flights booked')
        
        # Hotel
        doc.add_heading('Hotel', level=1)
        if hotel:
            doc.add_paragraph(f'Name: {hotel.get("name", "Unknown")}')
            if hotel.get('rating'):
                doc.add_paragraph(f'Rating: {hotel.get("rating")}⭐')
            doc.add_paragraph(f'Cost: ${hotel.get("total", 0):.2f}')
        else:
            doc.add_paragraph('No hotel booked')
        
        # Events
        if events_data and events_data.get("events"):
            doc.add_heading('Local Events', level=1)
            for event in events_data.get("events", []):
                doc.add_paragraph(
                    f'• {event.get("title")} - {event.get("start_date")}',
                    style='List Bullet'
                )
        
        # Weather
        if weather_info:
            doc.add_heading('Weather', level=1)
            doc.add_paragraph(f'Temperature: {weather_info.get("temperature", "N/A")}°F')
            doc.add_paragraph(f'Conditions: {weather_info.get("description", "N/A")}')
        
        # Save to bytes
        from io import BytesIO
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()
        
    except ImportError:
        # Fallback to text if python-docx not installed
        content = generate_trip_summary_text(trip_result, params, events_data, weather_info)
        return content.encode()

def execute_functions(function_calls: list, current_params: dict, trip_result: dict) -> dict:
    """Execute multiple function calls and return updates."""
    all_updates = {}
    all_results = []
    should_replan = False
    find_better = False
    current_hotel = None
    
    for func_call in function_calls:
        function_name = func_call["function"]
        arguments = func_call["arguments"]
        
        if function_name == "adjust_budget":
            new_budget = arguments.get("new_budget")
            if new_budget:
                all_updates["total_budget"] = float(new_budget)
                all_results.append(f"BUDGET_CHANGED:${new_budget}")
                should_replan = True
        
        elif function_name == "change_strategy":
            new_strategy = arguments.get("new_strategy")
            if new_strategy:
                all_updates["strategy"] = new_strategy
                strategy_display = {
                    'cheapest_overall': 'Cheapest Overall',
                    'splurge_flight': 'Splurge on Flights',
                    'splurge_hotel': 'Splurge on Hotel'
                }.get(new_strategy, new_strategy)
                all_results.append(f"STRATEGY_CHANGED:{strategy_display}")
                should_replan = True
        
        elif function_name == "adjust_iterations":
            new_iterations = arguments.get("new_iterations")
            if new_iterations:
                all_updates["max_iterations"] = int(new_iterations)
                all_results.append(f"ITERATIONS_CHANGED:{new_iterations}")
                should_replan = True
        
        elif function_name == "update_dates":
            departure = arguments.get("departure_date")
            return_date = arguments.get("return_date")
            if departure:
                all_updates["departure_date"] = departure
            if return_date:
                all_updates["return_date"] = return_date
            all_results.append(f"DATES_CHANGED:{departure} to {return_date}")
            should_replan = True
        
        elif function_name == "update_preferences":
            prefer_red_eyes = arguments.get("prefer_red_eyes")
            if prefer_red_eyes is not None:
                all_updates["prefer_red_eyes"] = prefer_red_eyes
                status = "enabled" if prefer_red_eyes else "disabled"
                all_results.append(f"REDEYE_{status.upper()}")
                should_replan = True
        
        elif function_name == "find_better_hotel":
            preference = arguments.get("preference", "any")
            find_better = True
            current_hotel = trip_result.get("data", {}).get("hotel", {})
            current_rating = current_hotel.get("rating", "Unknown")
            current_hotel["upgradePreference"] = preference
            all_results.append(f"FIND_BETTER_HOTEL:{preference}: current hotel ${current_hotel.get('total', 0)}")
            should_replan = True
        
        elif function_name == "search_events":
            preferences = arguments.get("preferences", [])
            
            # Search for events (returns raw data only)
            events_result = search_events(
                destination=st.session_state.current_params['destination'],
                start_date=st.session_state.current_params['departure_date'],
                end_date=st.session_state.current_params['return_date'],
                preferences=preferences,
                exclude_ids=st.session_state.excluded_event_ids,
                max_events=8,
            )
            
            # Store events for panel display
            if events_result["has_events"]:
                st.session_state.current_events = events_result["panel_html"]
                st.session_state.current_events_data = events_result
                st.session_state.excluded_event_ids = events_result["event_ids"]
            
            # Return RAW DATA - LLM will generate response
            all_results.append(json.dumps(events_result))
            should_replan = False
        
        elif function_name in ["search_flights", "search_hotels"]:
            should_replan = True
            all_results.append(f"SEARCH_NEW_{function_name.upper()}")
    
    return {
        "updates": all_updates,
        "results": all_results,
        "should_replan": should_replan,
        "find_better": find_better,
        "current_hotel": current_hotel
    }


def get_conversational_response(user_message: str, trip_data: dict, chat_history: list, function_results: list = None) -> str:
    """Get a natural language response from the LLM."""
    
    # Build rich context
    flights = trip_data.get("data", {}).get("flights", [])
    hotel = trip_data.get("data", {}).get("hotel", {})
    metadata = trip_data.get("metadata", {})
    
    # Format flight information nicely
    flight_info = []
    red_eye_count = 0
    for i, f in enumerate(flights, 1):
        direction = "Outbound" if i == 1 else "Return"
        airline = f.get('airline', 'Unknown')
        airline_code = f.get('airline_code', '')
        if airline_code and airline != airline_code:
            airline_display = f"{airline} ({airline_code})"
        else:
            airline_display = airline
        
        # Check if red-eye
        dep = f.get('departure_date', '')
        is_red_eye = False
        if " " in dep:
            try:
                t = datetime.strptime(dep.split(" ")[1][:5], "%H:%M")
                hour = t.hour + t.minute / 60
                is_red_eye = hour >= 21 or hour < 5
                if is_red_eye:
                    red_eye_count += 1
            except:
                pass
        
        red_eye_marker = " 🌙" if is_red_eye else ""
        
        flight_info.append(f"{direction} Flight{red_eye_marker}: {airline_display} {f.get('flight_number', '')}")
        flight_info.append(f"  • Depart: {f.get('departure_date', 'Unknown')}")
        flight_info.append(f"  • Arrive: {f.get('arrival_date', 'Unknown')}")
        flight_info.append(f"  • Duration: {f.get('duration', 'Unknown')}")
        flight_info.append(f"  • Cost: ${f.get('cost', 0):.2f}")
    
    # Format hotel information
    hotel_info = []
    if hotel:
        hotel_info.append(f"Hotel: {hotel.get('name', 'Unknown')}")
        hotel_info.append(f"  • Total Cost: ${hotel.get('total', 0):.2f}")
        if hotel.get('rating'):
            hotel_info.append(f"  • Rating: {hotel.get('rating')}⭐")
    else:
        hotel_info.append("No hotel selected yet.")
    
    # Add weather if available
    weather_info = None
    if 'weather_cache' in st.session_state:
        dest = metadata.get('destination', '')
        weather_info = st.session_state.weather_cache.get(dest)
    
    # Parse any event data from function results
    event_data = None
    if function_results:
        for result in function_results:
            try:
                # Try to parse as JSON (event data)
                parsed = json.loads(result)
                if isinstance(parsed, dict) and "has_events" in parsed:
                    event_data = parsed
                    break
            except:
                pass
    
    context = {
        "current_trip": {
            "origin": metadata.get("origin", "Unknown"),
            "destination": metadata.get("destination", "Unknown"),
            "dates": f"{metadata.get('departure_date')} to {metadata.get('return_date')}",
            "total_budget": f"${metadata.get('total_budget', 0):.2f}",
            "total_cost": f"${trip_data.get('data', {}).get('total_cost', 0):.2f}",
            "remaining_budget": f"${trip_data.get('data', {}).get('remaining_budget', 0):.2f}",
            "strategy": metadata.get('strategy', 'cheapest_overall').replace('_', ' ').title(),
            "optimization_status": trip_data.get('data', {}).get('status', 'unknown'),
            "iterations_used": f"{trip_data.get('data', {}).get('iterations_used', 0)}/{metadata.get('max_iterations', 5)}",
            "red_eye_enabled": metadata.get('prefer_red_eyes', False),
            "red_eye_flights_found": red_eye_count,
            "weather": weather_info
        },
        "flights": "\n".join(flight_info) if flight_info else "No flights found.",
        "hotel": "\n".join(hotel_info) if hotel_info else "No hotel found.",
        "events": event_data
    }
    
    system_prompt = f"""You are a friendly and knowledgeable travel assistant. Your role is to help users plan their trip by having natural conversations.

CURRENT TRIP CONTEXT:
{json.dumps(context, indent=2)}

WEATHER INFORMATION:
{json.dumps(weather_info, indent=2) if weather_info else "No weather data available"}

RECENT FUNCTION RESULTS (including event data if any):
{json.dumps(function_results, indent=2) if function_results else "No recent changes"}

You have access to these functions:
- search_flights: Find new flight options
- search_hotels: Find new hotel options
- adjust_budget: Change total budget
- change_strategy: Change allocation strategy
- adjust_iterations: Change optimization attempts
- update_dates: Change travel dates
- update_preferences: Enable/disable red-eye flights
- find_better_hotel: Find a cheaper, more expensive, higher rated, or different hotel
- search_events: Find local events during your trip

Guidelines for your responses:
1. Be conversational and friendly - use emojis appropriately
2. When users ask questions, provide helpful answers using the context
3. If they ask for something you don't have, explain what you can do instead
4. After making changes, explain what happened in a natural way
5. Keep responses concise but warm - 2-3 sentences usually
6. If they're asking about flight times, reference the specific flights
7. If weather info is available, mention it when relevant
8. If red-eye is enabled and flights found, mention you found overnight options
9. If no red-eye flights available, suggest they try different dates
10. When finding better hotels:
    - If UPGRADE_SUCCESS: Be excited! Mention the rating improvement and ask what you can help with next.
    - If FOUND_ALTERNATIVE: Mention you found another option with same rating to stay within budget
    - If NO_BETTER_HOTEL: Explain no better options found and suggest increasing budget
    - Don't just say "searching" - respond to what actually happened which already presented to the user.
    
11. WHEN YOU SEARCH FOR EVENTS:
    - You'll receive data about what was found. Use this to craft natural responses.

    IF EVENTS WERE FOUND:
    - Mention the total number of events found
    - Highlight 1-2 interesting events (title, day, category)
    - Reference that they can see all events in the panel above
    - Ask if they want more like a specific type or different categories

    Example tone (NOT hardcoded - just examples):
    - "Great news! I found 8 events in New York! There's a Knicks game on Wednesday and Hamilton on Thursday - both look amazing! Check out the panel above for all the details. Want me to find more concerts like the jazz one?"
    - "Found 4 food events during your stay! The Taste of NYC festival on Tuesday has 100+ vendors. Take a look at the panel - any of these catch your eye?"
    - "I searched for music events and found 6 shows! The Blue Note has jazz on Friday night that's really popular. Want to see more options or try a different category?"

    IF NO EVENTS FOUND:
    - Be helpful and suggest trying different preferences or dates
    - Example: "I couldn't find any events matching '{{preferences}}' in {destination}. Want to try different dates or maybe explore food festivals instead?"

    IF USER ASKS FOR "MORE LIKE THAT":
    - They want similar events (same category)
    - You can call search_events again with refined preferences

    IF USER ASKS FOR "DIFFERENT" EVENTS:
    - They want different categories than what they saw
    - Call search_events with new preferences

    WHEN DATES OR FLIGHTS OR HOTELS CHANGE:
    - If user previously asked for events, ASK if they want to search again 
    - Example: "I've updated your dates. Would you like me to search for events during your new stay dates?"
    - Do NOT assume they want events automatically

Remember: You're helping a real person plan their vacation. Be excited for them! Be enthusiastic, reference the panel, and always offer next steps naturally!"""
    
    messages = [
        {"role": "system", "content": system_prompt},
        *[{"role": m["role"], "content": m["message"]} for m in chat_history[-8:]],
    ]
    
    messages.append({"role": "user", "content": user_message})
    
    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=messages,
            max_tokens=300,
            temperature=0.7,
            timeout=10
        )
        
        return response.choices[0].message.content
        
    except Exception as e:
        return f"I'm having trouble responding right now. Error: {str(e)}"


def process_chat_message(user_message: str, trip_data: dict, chat_history: list, current_params: dict) -> dict:
    """Process chat message - determine if function calls are needed."""
    
    # Build context for function calling
    context = {
        "current_trip": {
            "origin": trip_data.get("metadata", {}).get("origin", "Unknown"),
            "destination": trip_data.get("metadata", {}).get("destination", "Unknown"),
            "departure_date": trip_data.get("metadata", {}).get("departure_date"),
            "return_date": trip_data.get("metadata", {}).get("return_date"),
            "total_budget": trip_data.get("metadata", {}).get("total_budget", 0),
            "strategy": trip_data.get("metadata", {}).get("strategy", "cheapest_overall"),
            "total_cost": trip_data.get("data", {}).get("total_cost", 0),
            "iterations_used": trip_data.get("data", {}).get("iterations_used", 0),
            "max_iterations": trip_data.get("metadata", {}).get("max_iterations", 5),
            "prefer_red_eyes": current_params.get("prefer_red_eyes", False)
        }
    }
    
    system_prompt = f"""You are a travel assistant that helps users by calling functions when they want to make changes.

Current trip context:
{json.dumps(context, indent=2)}

Available functions:
- search_flights: When user wants to search for new flights
- search_hotels: When user wants to search for new hotels  
- adjust_budget: When user wants to change budget
- change_strategy: When user wants to change allocation strategy
- adjust_iterations: When user wants to change optimization iterations
- update_dates: When user wants to change travel dates
- update_preferences: When user wants to enable/disable red-eye flights
- find_better_hotel: When user says 'find better hotel', 'cheaper hotel', 'more expensive hotel', 'different hotel', 'another hotel', or 'upgrade hotel'
- search_events: When user wants to find local events, things to do, concerts, etc.

If the user is just asking a question or having a conversation, do NOT call any functions.
Only call functions when they explicitly want to make changes to the trip, search for events, or update their preferences.
Do not call search_hotels, search_flights, or find_better_hotel for general questions like summaries, comparisons, explanations, or clarifications.
You can call multiple functions if the user makes multiple requests in one message."""

    messages = [
        {"role": "system", "content": system_prompt},
        *[{"role": m["role"], "content": m["message"]} for m in chat_history[-5:]],
        {"role": "user", "content": user_message}
    ]
    
    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=messages,
            tools=CHAT_FUNCTIONS,
            tool_choice="auto",
            max_tokens=200,
            temperature=0.0,
            timeout=10
        )
        
        message = response.choices[0].message
        
        if message.tool_calls and len(message.tool_calls) > 0:
            # Extract function calls
            function_calls = []
            for tool_call in message.tool_calls:
                function_calls.append({
                    "function": tool_call.function.name,
                    "arguments": json.loads(tool_call.function.arguments)
                })

            # Guard against over-eager tool calls for normal conversation.
            guarded_calls = []
            for call in function_calls:
                fn = call["function"]
                if fn == "find_better_hotel" and not should_find_better_hotel(user_message):
                    continue
                if fn == "search_hotels" and not should_search_hotels(user_message):
                    continue
                if fn == "search_flights" and not should_search_flights(user_message):
                    continue
                guarded_calls.append(call)

            if guarded_calls:
                return {
                    "type": "function_calls",
                    "function_calls": guarded_calls
                }
            return {
                "type": "conversation",
                "function_calls": []
            }
        else:
            return {
                "type": "conversation",
                "function_calls": []
            }
            
    except Exception as e:
        print(f"Error in process_chat_message: {e}")
        return {
            "type": "conversation",
            "function_calls": []
        }


# ==========================================================
# SESSION STATE
# ==========================================================

if 'trip_result' not in st.session_state:
    st.session_state.trip_result = None
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'optimization_iterations' not in st.session_state:
    st.session_state.optimization_iterations = []
if 'weather_cache' not in st.session_state:
    st.session_state.weather_cache = {}
if 'current_events' not in st.session_state:
    st.session_state.current_events = None
if 'current_events_data' not in st.session_state:
    st.session_state.current_events_data = None
if 'excluded_event_ids' not in st.session_state:
    st.session_state.excluded_event_ids = []
if 'trip_history' not in st.session_state:
    st.session_state.trip_history = []
if 'current_trip_id' not in st.session_state:
    st.session_state.current_trip_id = -1
if 'trip_version' not in st.session_state:
    st.session_state.trip_version = 0
if 'current_params' not in st.session_state:
    st.session_state.current_params = {
        'origin': 'San Francisco',
        'destination': 'New York City',
        'departure_date': (datetime.now() + timedelta(days=30)).strftime("%Y-%m-%d"),
        'return_date': (datetime.now() + timedelta(days=37)).strftime("%Y-%m-%d"),
        'total_budget': 1500.0,
        'strategy': 'cheapest_overall',
        'prefer_red_eyes': False,
        'max_iterations': 5,
    }
if 'sidebar_origin' not in st.session_state:
    st.session_state.sidebar_origin = st.session_state.current_params['origin']
if 'sidebar_destination' not in st.session_state:
    st.session_state.sidebar_destination = st.session_state.current_params['destination']
if 'sidebar_departure' not in st.session_state:
    st.session_state.sidebar_departure = datetime.strptime(st.session_state.current_params['departure_date'], "%Y-%m-%d").date()
if 'sidebar_return_date' not in st.session_state:
    st.session_state.sidebar_return_date = datetime.strptime(st.session_state.current_params['return_date'], "%Y-%m-%d").date()
if 'sidebar_total_budget' not in st.session_state:
    st.session_state.sidebar_total_budget = float(st.session_state.current_params['total_budget'])
if 'sidebar_strategy' not in st.session_state:
    st.session_state.sidebar_strategy = st.session_state.current_params['strategy']
if 'sidebar_prefer_red_eyes' not in st.session_state:
    st.session_state.sidebar_prefer_red_eyes = st.session_state.current_params['prefer_red_eyes']
if 'sidebar_max_iterations' not in st.session_state:
    st.session_state.sidebar_max_iterations = st.session_state.current_params['max_iterations']

if 'last_chat_time' not in st.session_state:
    st.session_state.last_chat_time = 0
if 'show_debug' not in st.session_state:
    st.session_state.show_debug = False
if 'processing_message' not in st.session_state:
    st.session_state.processing_message = False
if 'show_save_dialog' not in st.session_state:
    st.session_state.show_save_dialog = False
if 'save_format' not in st.session_state:
    st.session_state.save_format = "Text (.txt)"


# ==========================================================
# UI HEADER
# ==========================================================

st.markdown("""
<div class="main-header">
    <h1>✈️ AI TRAVEL PLANNER</h1>
    <p>Your Personal AI Travel Assistant ✈️ Flights 🏨 Hotels 🎉 Events 🌤️ Weather 💬 Chat</p>
</div>
""", unsafe_allow_html=True)


# ==========================================================
# SIDEBAR - Trip Parameters & History (removed adults)
# ==========================================================

with st.sidebar:
    st.markdown("### 🎯 Your Trip Details")
    
    origin = st.text_input("From", key="sidebar_origin")
    destination = st.text_input("To", key="sidebar_destination")
    
    col1, col2 = st.columns(2)
    with col1:
        departure = st.date_input(
            "Depart",
            key="sidebar_departure"
        )
    with col2:
        return_date = st.date_input(
            "Return",
            key="sidebar_return_date"
        )
    
    if return_date <= departure:
        st.error("Return date must be after departure date!")
        valid_dates = False
    else:
        valid_dates = True
        nights = (return_date - departure).days
        st.caption(f"📅 {nights} nights")
    
    total_budget = st.number_input(
        "Budget ($)",
        min_value=100.0,
        max_value=10000.0,
        step=100.0,
        key="sidebar_total_budget"
    )
    
    strategy_options = {
        "cheapest_overall": "Cheapest Overall",
        "splurge_flight": "Splurge on Flights",
        "splurge_hotel": "Splurge on Hotel"
    }
    strategy = st.selectbox(
        "Strategy",
        options=list(strategy_options.keys()),
        format_func=lambda x: strategy_options[x],
        key="sidebar_strategy"
    )
    
    col3, col4 = st.columns(2)
    with col3:
        prefer_red_eyes = st.checkbox(
            "🌙 Prefer Red-Eye",
            key="sidebar_prefer_red_eyes"
        )
    
    max_iterations = st.slider(
        "Optimization Attempts",
        min_value=1,
        max_value=20,
        key="sidebar_max_iterations"
    )
    
    show_debug = st.checkbox("🔧 Show Debug Info", value=st.session_state.show_debug)
    st.session_state.show_debug = show_debug
    
    # Plan New Trip Button
    if st.button("🔍 Plan New Trip", use_container_width=True, disabled=not valid_dates):
        st.session_state.current_params.update({
            'origin': st.session_state.sidebar_origin,
            'destination': st.session_state.sidebar_destination,
            'departure_date': st.session_state.sidebar_departure.strftime("%Y-%m-%d"),
            'return_date': st.session_state.sidebar_return_date.strftime("%Y-%m-%d"),
            'total_budget': float(st.session_state.sidebar_total_budget),
            'strategy': st.session_state.sidebar_strategy,
            'prefer_red_eyes': st.session_state.sidebar_prefer_red_eyes,
            'max_iterations': st.session_state.sidebar_max_iterations,
        })
        
        # Clear events when planning new trip
        st.session_state.current_events = None
        st.session_state.current_events_data = None
        st.session_state.excluded_event_ids = []
        
        with st.spinner("Planning your perfect trip..."):
            try:
                result = plan_trip(
                    origin=st.session_state.sidebar_origin,
                    destination=st.session_state.sidebar_destination,
                    departure_date=st.session_state.sidebar_departure.strftime("%Y-%m-%d"),
                    return_date=st.session_state.sidebar_return_date.strftime("%Y-%m-%d"),
                    total_budget=float(st.session_state.sidebar_total_budget),
                    strategy=Strategy(st.session_state.sidebar_strategy),
                    prefer_red_eyes=st.session_state.sidebar_prefer_red_eyes,
                    max_iterations=st.session_state.sidebar_max_iterations,
                )
                
                st.session_state.trip_result = result
                st.session_state.optimization_iterations = result.get("optimization_history", [])
                
                # Clear old weather and fetch new
                st.session_state.weather_cache = {}
                
                # Get a warm welcome with trip context
                trip_context = f"I just planned a trip from {st.session_state.sidebar_origin} to {st.session_state.sidebar_destination} from {st.session_state.sidebar_departure.strftime('%B %d')} to {st.session_state.sidebar_return_date.strftime('%B %d')} with a budget of ${st.session_state.sidebar_total_budget}. Can you welcome me and help me with my trip?"

                welcome_message = get_conversational_response(
                    trip_context,
                    result,
                    [],
                    None
                )
                st.session_state.chat_history = [{"role": "assistant", "message": welcome_message}]
                
                # Save to history (auto-save)
                save_trip_to_history()
                
                st.rerun()
                
            except Exception as e:
                st.error(f"Error: {str(e)}")
    
    st.markdown("---")
    
    # Current Trip Summary
    if st.session_state.trip_result:
        total = st.session_state.trip_result.get("data", {}).get("total_cost", 0)
        st.metric("Current Total", f"${total:.2f}")
    
    # ==========================================================
    # TRIP HISTORY SECTION
    # ==========================================================
    st.markdown("### 📜 Trip History")
    
    # New Trip Button
    if st.button("➕ Start New Trip", use_container_width=True):
        start_new_trip()
    
    # Display trip history
    if st.session_state.trip_history:
        st.markdown("---")
        for trip in st.session_state.trip_history[:5]:  # Show last 5 trips
            # Determine if this is the current trip
            is_current = trip["id"] == st.session_state.current_trip_id
            
            # Create columns for trip display and delete button
            col1, col2 = st.columns([0.9, 0.1])
            
            with col1:
                # Clickable trip card
                if st.button(
                    f"**{trip['origin']} → {trip['destination']}**\n\n{trip['dates']}\n\n💰 ${trip['total_cost']:.0f}",
                    key=f"trip_{trip['id']}",
                    use_container_width=True,
                    help="Click to load this trip"
                ):
                    load_trip_from_history(trip["id"])
            
            with col2:
                # Delete button
                if st.button("X ", key=f"del_{trip['id']}", help="Delete this trip"):
                    delete_trip_from_history(trip["id"])
    else:
        st.caption("No past trips yet")


# ==========================================================
# MAIN CONTENT - Results and Chat
# ==========================================================

col_main, col_chat = st.columns([0.7, 0.3])

with col_main:
    if st.session_state.trip_result:
        result = st.session_state.trip_result
        data = result.get("data", {})
        metadata = result.get("metadata", {})
        
        # Metrics
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Total Cost", f"${data.get('total_cost', 0):.2f}")
        with col2:
            st.metric("Remaining", f"${data.get('remaining_budget', 0):.2f}")
        with col3:
            if data.get('flight_ratio'):
                st.metric("Flights %", f"{data['flight_ratio']*100:.0f}%")
            else:
                iterations_used = data.get("iterations_used", 0)
                max_iterations = metadata.get("max_iterations", 5)
                st.metric("Iterations", f"{iterations_used}/{max_iterations}")
        with col4:
            if data.get('hotel_ratio'):
                st.metric("Hotel %", f"{data['hotel_ratio']*100:.0f}%")
            else:
                status = data.get("status", "unknown")
                st.metric("Status", status.replace("_", " ").title())
        
        # Red-eye indicator
        if metadata.get('prefer_red_eyes', False):
            st.markdown('<span class="red-eye-badge">🌙 Red-eye mode enabled</span>', unsafe_allow_html=True)
        
        # ==========================================================
        # SAVE BUTTON AND DIALOG
        # ==========================================================
        col_save1, col_save2, col_save3 = st.columns([1, 2, 1])
        with col_save2:
            if st.button("💾 Save Trip", use_container_width=True, key="show_save_dialog_btn"):
                st.session_state.show_save_dialog = True
                st.rerun()
        
        # Save Dialog
        if st.session_state.show_save_dialog:
            with st.container():
                st.markdown("---")
                st.markdown("### 💾 Save Trip")
                
                col_f1, col_f2 = st.columns(2)
                with col_f1:
                    # Format selection
                    format_options = ["Text (.txt)", "Excel (.xlsx)", "Word (.docx)", "CSV (.csv)"]
                    format_idx = format_options.index(st.session_state.save_format) if st.session_state.save_format in format_options else 0
                    save_format = st.selectbox("Select format", format_options, index=format_idx, key="format_select")
                    st.session_state.save_format = save_format
                
                with col_f2:
                    # Filename
                    default_name = f"Trip_{metadata.get('origin', 'Trip')}_to_{metadata.get('destination', 'Dest')}"
                    filename = st.text_input("Filename", value=default_name, key="filename_input")
                
                # Prepare data for download based on current selections
                events_data = st.session_state.current_events_data
                weather_info = None
                if 'weather_cache' in st.session_state and metadata.get('destination') in st.session_state.weather_cache:
                    weather_info = st.session_state.weather_cache[metadata.get('destination')]
                
                # Generate file data based on selected format
                file_data = None
                mime_type = None
                file_ext = ""
                
                if save_format == "Text (.txt)":
                    content = generate_trip_summary_text(result, st.session_state.current_params, events_data, weather_info)
                    file_data = content.encode()
                    mime_type = "text/plain"
                    file_ext = "txt"
                elif save_format == "Excel (.xlsx)":
                    excel_data = generate_excel_export(result, st.session_state.current_params, events_data, weather_info)
                    file_data = excel_data.getvalue()
                    mime_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    file_ext = "xlsx"
                elif save_format == "Word (.docx)":
                    file_data = generate_word_export(result, st.session_state.current_params, events_data, weather_info)
                    mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    file_ext = "docx"
                elif save_format == "CSV (.csv)":
                    content = generate_csv_export(result, st.session_state.current_params, events_data, weather_info)
                    file_data = content.encode()
                    mime_type = "text/csv"
                    file_ext = "csv"
                
                # Clean filename
                safe_filename = filename.replace(' ', '_').replace('/', '_').replace('\\', '_')
                if not safe_filename:
                    safe_filename = f"Trip_{metadata.get('origin', 'Trip')}_to_{metadata.get('destination', 'Dest')}"
                
                full_filename = f"{safe_filename}.{file_ext}"
                
                col_b1, col_b2, col_b3 = st.columns([1, 1, 1])
                with col_b1:
                    # Download button
                    st.download_button(
                        label="📥 Download",
                        data=file_data,
                        file_name=full_filename,
                        mime=mime_type,
                        key="download_file",
                        use_container_width=True
                    )
                
                with col_b2:
                    if st.button("Close", use_container_width=True, key="close_dialog"):
                        st.session_state.show_save_dialog = False
                        st.rerun()
                
                st.markdown("---")
        
        # Flights
        st.markdown("### ✈️ Your Flights")
        flights = data.get('flights', [])
        if flights:
            red_eye_count = 0
            for i, flight in enumerate(flights, 1):
                airline = flight.get('airline', 'Unknown')
                airline_code = flight.get('airline_code', '')
                flight_num = flight.get('flight_number', '')
                dep = flight.get('departure_date', '')
                
                # Check if red-eye
                is_red_eye = False
                if " " in dep:
                    try:
                        t = datetime.strptime(dep.split(" ")[1][:5], "%H:%M")
                        hour = t.hour + t.minute / 60
                        is_red_eye = hour >= 21 or hour < 5
                        if is_red_eye:
                            red_eye_count += 1
                    except:
                        pass
                
                if airline_code and airline != airline_code:
                    airline_display = f"{airline} ({airline_code})"
                else:
                    airline_display = airline
                
                red_eye_badge = " 🌙" if is_red_eye else ""
                
                with st.container():
                    cols = st.columns([2, 1, 1, 1])
                    with cols[0]:
                        st.markdown(f"**{airline_display}{red_eye_badge}**")
                        st.caption(f"{flight_num} • {flight.get('home_airport')} → {flight.get('destination')}")
                    with cols[1]:
                        st.markdown(f"Depart\n{flight.get('departure_date', 'N/A')}")
                    with cols[2]:
                        st.markdown(f"Arrive\n{flight.get('arrival_date', 'N/A')}")
                    with cols[3]:
                        st.markdown(f"**${flight.get('cost', 0):.2f}**")
                    st.divider()
            
            if metadata.get('prefer_red_eyes', False) and red_eye_count == 0:
                st.warning("No red-eye flights available for this route. Try different dates!")
        else:
            st.info("No flights found for this route")
        
        # Hotel
        st.markdown("### 🏨 Your Hotel")
        hotel = data.get('hotel')
        hotel_rankings = metadata.get('hotel_rankings', {}) if metadata else {}
        if hotel:
            cols = st.columns([3, 1, 1])
            with cols[0]:
                st.markdown(f"**{hotel.get('name', 'Unknown')}**")
                rating = hotel.get('rating')
                if rating:
                    st.caption(f"Rating: {rating}⭐")

                badges = []
                if hotel.get('isCheapest'):
                    badges.append("💸 Cheapest")
                if hotel.get('isBestValue'):
                    badges.append("⭐ Best value")
                if hotel.get('isBestRated'):
                    badges.append("🏆 Best rated")
                if hotel.get('priceRank'):
                    badges.append(f"Price rank #{hotel.get('priceRank')}")
                if hotel.get('overBudget'):
                    badges.append(f"Over budget by ${hotel.get('overBudget'):.2f}")
                if badges:
                    st.caption(" • ".join(badges))
            with cols[1]:
                st.markdown("Total")
            with cols[2]:
                st.markdown(f"**${hotel.get('total', 0):.2f}**")

            ranking_cards = []
            for label, key in [("Cheapest", "cheapest"), ("Best value", "best_value"), ("Best rated", "best_rating")]:
                ranked = hotel_rankings.get(key)
                if ranked and ranked.get('hotelId') != hotel.get('hotelId'):
                    ranking_cards.append((label, ranked))

            if ranking_cards:
                st.markdown("#### Other good hotel picks")
                cols = st.columns(len(ranking_cards))
                for col, (label, ranked) in zip(cols, ranking_cards):
                    with col:
                        st.markdown(f"**{label}**")
                        st.caption(ranked.get('name', 'Unknown'))
                        rating = ranked.get('rating')
                        total = float(ranked.get('total', 0) or 0)
                        extras = []
                        if rating:
                            extras.append(f"{rating}⭐")
                        extras.append(f"${total:.2f}")
                        st.write(" • ".join(extras))
        else:
            st.info("No hotel found")
        
        # Events Panel
        if st.session_state.current_events:
            st.markdown("### 🎉 Local Events")
            st.markdown(st.session_state.current_events, unsafe_allow_html=True)
        
        # Weather
        if hotel or flights:
            st.markdown("### 🌤️ Destination Weather")
            with st.spinner("Checking forecast..."):
                dest_city = metadata.get('destination', 'New York City')
                arrival_date = metadata.get('departure_date')
                
                city_for_weather = dest_city.split(',')[0].strip()
                
                if dest_city not in st.session_state.weather_cache:
                    weather = get_weather_forecast(city_for_weather, arrival_date)
                    if weather:
                        st.session_state.weather_cache[dest_city] = weather
                else:
                    weather = st.session_state.weather_cache[dest_city]
                
                if weather:
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.metric("Temperature", f"{weather['temperature']}°F", 
                                 help="Forecast for arrival day")
                    with col2:
                        st.metric("Conditions", f"{weather['emoji']} {weather['description']}")
                    with col3:
                        st.metric("Wind", f"{weather['wind']} mph")
                    
                    if weather.get('humidity'):
                        st.caption(f"Humidity: {weather['humidity']}%")
                else:
                    st.caption("🌡️ Weather forecast unavailable")
        
        # Debug section
        if st.session_state.show_debug:
            with st.expander("🔧 Optimization Details"):
                iterations_used = data.get("iterations_used", 0)
                max_iterations = metadata.get("max_iterations", 0)
                
                col1, col2 = st.columns(2)
                with col1:
                    st.metric("Iterations Used", f"{iterations_used}/{max_iterations}")
                with col2:
                    status = data.get("status", "unknown")
                    st.metric("Status", status.replace("_", " ").title())
                
                history = result.get("optimization_history", [])
                if history:
                    st.write(f"**Optimization Steps:**")
                    for i, it in enumerate(history):
                        st.json(it)
    else:
        st.info("👈 Enter your trip details and click 'Plan New Trip' to get started!")

with col_chat:
    st.markdown("### 💬 Chat with Your Assistant")
    
    chat_container = st.container(height=500)
    with chat_container:
        if st.session_state.chat_history:
            for msg in st.session_state.chat_history:
                if msg["role"] == "user":
                    st.markdown(f'<div class="chat-message-user">👤 {msg["message"]}</div>', unsafe_allow_html=True)
                else:
                    st.markdown(f'<div class="chat-message-bot">🤖 {msg["message"]}</div>', unsafe_allow_html=True)
        else:
            st.caption("Plan a trip first, then chat with me!")
    
    if st.session_state.trip_result and not st.session_state.processing_message:
        user_message = st.chat_input("Ask me anything about your trip...")
        
        current_time = time.time()
        if user_message and (current_time - st.session_state.last_chat_time) >= RATE_LIMIT_SECONDS:
            st.session_state.last_chat_time = current_time
            st.session_state.processing_message = True
            
            # Add user message immediately
            st.session_state.chat_history.append({"role": "user", "message": user_message})
            
            with st.spinner("🤔 Thinking..."):
                # Step 1: Determine if function calls are needed
                process_result = process_chat_message(
                    user_message,
                    st.session_state.trip_result,
                    st.session_state.chat_history[:-1],
                    st.session_state.current_params
                )
                
                function_results = None
                should_replan = False
                find_better = False
                current_hotel_data = None
                current_flights_data = None
                
                # Step 2: Execute any function calls
                if process_result["type"] == "function_calls" and process_result["function_calls"]:
                    exec_result = execute_functions(
                        process_result["function_calls"],
                        st.session_state.current_params,
                        st.session_state.trip_result
                    )
                    
                    function_results = exec_result.get("results", [])
                    find_better = exec_result.get("find_better", False)
                    current_hotel_data = exec_result.get("current_hotel")
                    current_flights_data = st.session_state.trip_result.get("data", {}).get("flights", [])
                    should_replan = exec_result.get("should_replan", False)
                    
                    # Update params
                    if exec_result["updates"]:
                        for key, value in exec_result["updates"].items():
                            st.session_state.current_params[key] = value
                
                # Step 3: Replan if needed
                if should_replan:
                    with st.spinner("🔄 Updating your trip..."):
                        try:
                            new_result = plan_trip(
                                origin=st.session_state.current_params['origin'],
                                destination=st.session_state.current_params['destination'],
                                departure_date=st.session_state.current_params['departure_date'],
                                return_date=st.session_state.current_params['return_date'],
                                total_budget=float(st.session_state.current_params['total_budget']),
                                strategy=Strategy(st.session_state.current_params['strategy']),
                                prefer_red_eyes=st.session_state.current_params['prefer_red_eyes'],
                                max_iterations=st.session_state.current_params['max_iterations'],
                                find_better_hotel=find_better,
                                current_hotel=current_hotel_data if find_better else None,
                                current_flights=current_flights_data if find_better else None
                            )
                            
                            # Enhance function results with upgrade information
                            if find_better and current_hotel_data:
                                new_hotel = new_result.get("data", {}).get("hotel", {})
                                if new_hotel and new_hotel.get("hotelId") != current_hotel_data.get("hotelId"):
                                    old_price = float(current_hotel_data.get("total", 0) or 0)
                                    new_price = float(new_hotel.get("total", 0) or 0)
                                    if new_price < old_price:
                                        function_results.append(f"FOUND_CHEAPER_OPTION: Found a cheaper hotel at ${new_price:.2f} (was ${old_price:.2f})")
                                    elif new_price > old_price:
                                        function_results.append(f"FOUND_PRICIER_OPTION: Found a pricier hotel at ${new_price:.2f} (was ${old_price:.2f})")
                                    else:
                                        function_results.append("FOUND_ALTERNATIVE: Found another hotel option at a similar price")
                                else:
                                    function_results.append("NO_BETTER_HOTEL: Could not find better hotel within budget")
                            
                            st.session_state.trip_result = new_result
                            st.session_state.optimization_iterations = new_result.get("optimization_history", [])
                            
                            # Clear weather cache for new search
                            if 'weather_cache' in st.session_state:
                                st.session_state.weather_cache = {}
                            
                            # AUTO-SAVE after any update
                            save_trip_to_history()
                                
                        except Exception as e:
                            function_results = [f"Error updating: {str(e)}"]
                            print(f"Replan error: {e}")
                
                # Step 4: Get natural language response
                response = get_conversational_response(
                    user_message,
                    st.session_state.trip_result,
                    st.session_state.chat_history[:-1],
                    function_results
                )
                
                st.session_state.chat_history.append({"role": "assistant", "message": response})
                
                # AUTO-SAVE after chat
                save_trip_to_history()
                
                # Trim history
                if len(st.session_state.chat_history) > MAX_CHAT_HISTORY:
                    st.session_state.chat_history = st.session_state.chat_history[-MAX_CHAT_HISTORY:]
            
            st.session_state.processing_message = False
            st.rerun()
        
        elif user_message:
            wait_time = RATE_LIMIT_SECONDS - (current_time - st.session_state.last_chat_time)
            if wait_time > 0:
                st.warning(f"Please wait {wait_time:.1f}s")


# ==========================================================
# FOOTER
# ==========================================================

st.markdown("---")
st.markdown("""
<div style="text-align: center; color: #666; padding: 1rem;">
    <p>Your Personal AI Travel Assistant | Powered by Amadeus, OpenAI, OpenWeatherMap & PredictHQ</p>
</div>
""", unsafe_allow_html=True)